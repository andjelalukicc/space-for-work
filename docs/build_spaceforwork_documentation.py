from pathlib import Path

from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.opc.constants import RELATIONSHIP_TYPE
from docx.shared import Cm, Pt, RGBColor


DOCS_DIR = Path(__file__).resolve().parent
OUT = DOCS_DIR / "Space_For_Work_Projektna_Dokumentacija_2026.docx"
ASSET_DIR = DOCS_DIR / "assets"
DIAGRAM_DIR = DOCS_DIR / "figma-export" / "png"
REPO_URL = "https://github.com/andjelalukicc/space-for-work"
PORTAL_URL = "http://127.0.0.1:8888/spaceforwork-portal.html"
SCREENSHOT_DIR = ASSET_DIR
FONT = "Times New Roman"


def set_run(run, size=12, bold=False, italic=False, color=None):
    run.font.name = FONT
    run._element.rPr.rFonts.set(qn("w:eastAsia"), FONT)
    run.font.size = Pt(size)
    run.bold = bold
    run.italic = italic
    if color:
        run.font.color.rgb = color


def style_paragraph(paragraph, align=None, space_before=0, space_after=6, line_spacing=1.05):
    if align is not None:
        paragraph.alignment = align
    paragraph.paragraph_format.space_before = Pt(space_before)
    paragraph.paragraph_format.space_after = Pt(space_after)
    paragraph.paragraph_format.line_spacing = line_spacing


def add_paragraph(
    doc,
    text="",
    *,
    align=None,
    size=12,
    bold=False,
    italic=False,
    color=None,
    space_before=0,
    space_after=6,
    first_line=False,
):
    paragraph = doc.add_paragraph()
    style_paragraph(paragraph, align, space_before, space_after)
    if first_line:
        paragraph.paragraph_format.first_line_indent = Cm(0.6)
    parts = str(text).split("\n")
    for idx, part in enumerate(parts):
        if idx:
            paragraph.add_run().add_break()
        run = paragraph.add_run(part)
        set_run(run, size=size, bold=bold, italic=italic, color=color)
    return paragraph


def add_blank(doc, count=1):
    for _ in range(count):
        add_paragraph(doc, "", space_after=8)


def add_section_title(doc, number, title):
    paragraph = add_paragraph(
        doc,
        f"   {number}.        {title}",
        size=18,
        bold=True,
        space_before=10,
        space_after=12,
    )
    paragraph.paragraph_format.keep_with_next = True
    return paragraph


def add_subtitle(doc, number, title):
    paragraph = add_paragraph(
        doc,
        f"{number}. {title}",
        size=13,
        bold=True,
        space_before=10,
        space_after=8,
    )
    paragraph.paragraph_format.keep_with_next = True
    return paragraph


def add_minor_title(doc, number, title):
    paragraph = add_paragraph(
        doc,
        f"{number}. {title}",
        size=12,
        bold=True,
        space_before=8,
        space_after=6,
    )
    paragraph.paragraph_format.keep_with_next = True
    return paragraph


def add_bullets(doc, items):
    for item in items:
        paragraph = doc.add_paragraph(style="List Bullet")
        paragraph.paragraph_format.space_after = Pt(4)
        paragraph.paragraph_format.line_spacing = 1.05
        run = paragraph.add_run(item)
        set_run(run, size=12)


def add_numbered(doc, items):
    for item in items:
        paragraph = doc.add_paragraph(style="List Number")
        paragraph.paragraph_format.space_after = Pt(4)
        run = paragraph.add_run(item)
        set_run(run, size=12)


def set_no_borders(table):
    tbl = table._tbl
    tbl_pr = tbl.tblPr
    borders = tbl_pr.first_child_found_in("w:tblBorders")
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        tag = "w:" + edge
        element = borders.find(qn(tag))
        if element is None:
            element = OxmlElement(tag)
            borders.append(element)
        element.set(qn("w:val"), "nil")


def add_hyperlink(paragraph, text, url):
    part = paragraph.part
    rel_id = part.relate_to(url, RELATIONSHIP_TYPE.HYPERLINK, is_external=True)
    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), rel_id)
    run_element = OxmlElement("w:r")
    run_props = OxmlElement("w:rPr")
    color = OxmlElement("w:color")
    color.set(qn("w:val"), "0563C1")
    underline = OxmlElement("w:u")
    underline.set(qn("w:val"), "single")
    run_props.append(color)
    run_props.append(underline)
    run_element.append(run_props)
    text_element = OxmlElement("w:t")
    text_element.text = text
    run_element.append(text_element)
    hyperlink.append(run_element)
    paragraph._p.append(hyperlink)


def add_image(doc, image_path, width_cm, caption=None):
    image_path = Path(image_path)
    if not image_path.exists():
        add_paragraph(doc, f"[Slika nije pronadjena: {image_path.name}]", italic=True)
        return
    paragraph = doc.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = paragraph.add_run()
    run.add_picture(str(image_path), width=Cm(width_cm))
    if caption:
        caption_p = add_paragraph(
            doc,
            caption,
            align=WD_ALIGN_PARAGRAPH.CENTER,
            size=10,
            italic=True,
            color=RGBColor(96, 96, 96),
            space_after=8,
        )
        caption_p.paragraph_format.keep_with_next = False


def configure_document():
    doc = Document()
    section = doc.sections[0]
    section.page_width = Cm(21.59)
    section.page_height = Cm(27.94)
    section.top_margin = Cm(2.25)
    section.bottom_margin = Cm(2.0)
    section.left_margin = Cm(2.25)
    section.right_margin = Cm(2.25)

    normal = doc.styles["Normal"]
    normal.font.name = FONT
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), FONT)
    normal.font.size = Pt(12)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.05

    for style_name in ("List Bullet", "List Number"):
        style = doc.styles[style_name]
        style.font.name = FONT
        style._element.rPr.rFonts.set(qn("w:eastAsia"), FONT)
        style.font.size = Pt(12)

    return doc


def add_cover(doc):
    logo_table = doc.add_table(rows=1, cols=3)
    logo_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    set_no_borders(logo_table)
    logo_table.columns[0].width = Cm(4)
    logo_table.columns[1].width = Cm(10)
    logo_table.columns[2].width = Cm(4)

    left = logo_table.cell(0, 0)
    center = logo_table.cell(0, 1)
    right = logo_table.cell(0, 2)
    for cell in (left, center, right):
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.TOP

    if (ASSET_DIR / "uns-logo.png").exists():
        left.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.LEFT
        left.paragraphs[0].add_run().add_picture(str(ASSET_DIR / "uns-logo.png"), width=Cm(1.9))
    if (ASSET_DIR / "ftn-logo.png").exists():
        right.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT
        right.paragraphs[0].add_run().add_picture(str(ASSET_DIR / "ftn-logo.png"), width=Cm(1.9))

    center.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = center.paragraphs[0].add_run("UNIVERZITET U NOVOM SADU\nFAKULTET TEHNIČKIH NAUKA")
    set_run(r, size=13, bold=True)

    add_blank(doc, 5)
    add_paragraph(doc, "PROJEKTNA DOKUMENTACIJA", align=WD_ALIGN_PARAGRAPH.CENTER, size=20, bold=True)
    add_blank(doc, 3)
    add_paragraph(
        doc,
        "PREDMET: EKSPLOATACIJA, ODRŽAVANJE I NADOGRADNJA\nINFORMACIONIH SISTEMA",
        align=WD_ALIGN_PARAGRAPH.CENTER,
        size=13,
        bold=True,
        space_after=24,
    )
    add_paragraph(
        doc,
        "TEMA: SPACE FOR WORK - INFORMACIONI SISTEM ZA COWORKING PROSTOR",
        align=WD_ALIGN_PARAGRAPH.CENTER,
        size=15,
        bold=True,
    )
    add_blank(doc, 4)
    add_paragraph(
        doc,
        "GitHub repozitorijum i demo portal:",
        align=WD_ALIGN_PARAGRAPH.CENTER,
        size=11,
        bold=True,
        space_after=4,
    )
    link_p = doc.add_paragraph()
    link_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_hyperlink(link_p, REPO_URL, REPO_URL)
    portal_p = doc.add_paragraph()
    portal_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_hyperlink(portal_p, PORTAL_URL, PORTAL_URL)
    add_blank(doc, 4)

    info_table = doc.add_table(rows=1, cols=2)
    set_no_borders(info_table)
    info_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    left_info = info_table.cell(0, 0)
    right_info = info_table.cell(0, 1)
    left_info.width = Cm(8.2)
    right_info.width = Cm(8.2)

    left_info.text = ""
    p = left_info.paragraphs[0]
    for line in ("Profesor: Teodora Vučković", "Asistent: Anđela Todorić", "Asistent: Sara Kijanović"):
        if p.runs:
            p.add_run().add_break()
        run = p.add_run(line)
        set_run(run, size=12)

    right_info.text = ""
    p = right_info.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = p.add_run("Student: Anđela Lukić IT62/2022")
    set_run(run, size=12)

    doc.add_page_break()


def add_contents(doc):
    add_paragraph(doc, "Sadržaj:", size=16, bold=True, color=RGBColor(31, 78, 121), space_after=10)
    entries = [
        ("1.", "Opis realnog sistema", "3"),
        ("2.", "Korišćene tehnologije", "4"),
        ("3.", "UML DIJAGRAMI", "6"),
        ("4.", "Baza podataka", "10"),
        ("5.", "Opis predloženog rešenja", "12"),
        ("6.", "Zaključak", "18"),
    ]
    for number, title, page in entries:
        paragraph = doc.add_paragraph()
        style_paragraph(paragraph, space_after=3)
        tab_stops = paragraph.paragraph_format.tab_stops
        tab_stops.add_tab_stop(Cm(16.3))
        run = paragraph.add_run(f"{number}   {title}")
        set_run(run, size=12)
        dots = "." * max(12, 116 - len(title))
        run = paragraph.add_run(dots)
        set_run(run, size=9, color=RGBColor(90, 90, 90))
        run = paragraph.add_run(f"\t{page}")
        set_run(run, size=12)
    doc.add_page_break()


def add_real_system(doc):
    add_section_title(doc, "1", "Opis realnog sistema")
    add_paragraph(
        doc,
        "Space For Work je coworking prostor u Novom Sadu namenjen ljudima i timovima kojima treba "
        "radno okruženje koje mogu brzo da koriste, bez dugoročnih obaveza i dodatne organizacije. "
        "Prostor obuhvata više od 900 m2, 99 radnih mesta i 12 kancelarija, uz sale za sastanke, event "
        "prostor, open space zonu, private office opcije i Virtual Office uslugu.",
        first_line=True,
    )
    add_paragraph(
        doc,
        "Cilj sistema je da digitalizuje celokupno poslovanje coworking prostora: javna prezentacija "
        "usluga, prodaja paketa i članstva, upravljanje Virtual Office i private office ponudom, "
        "praćenje članova i transakcija, uz podršku za rezervaciju sala i phone booth-ova kada je "
        "potrebno. Time se izbegava ručno vođenje poruka, tabela i potvrda, a korisnik na jednom "
        "mestu dobija pregled ponude, paketa i statusa svojih zahteva.",
        first_line=True,
    )

    add_paragraph(doc, "Akteri:", bold=True, space_before=6)
    add_bullets(
        doc,
        [
            "Admin - zaposleni ili menadžer prostora koji se prijavljuje kroz administratorski panel i upravlja korisnicima, paketima, članstvima, Virtual Office zahtevima, transakcijama i resursima prostora.",
            "Korisnik - član, freelancer, remote radnik, startup ili tim koji pregleda coworking ponudu, kupuje paket, šalje Virtual Office zahtev i po potrebi rezerviše salu ili booth.",
            "Gost - posetilac koji može da pregleda javni deo sajta, vidi pakete i usluge, pošalje kontakt upit i započne proces kupovine ili obilaska prostora.",
            "Stripe - eksterni servis za testno plaćanje karticom i slanje webhook potvrde o uspešnoj uplati.",
        ],
    )

    add_paragraph(doc, "Funkcionisanje sistema:", bold=True, space_before=6)
    add_paragraph(
        doc,
        "Korisnik najpre dolazi na javni portal gde može da vidi coworking ponudu, membership pakete, "
        "Virtual Office, private office, galeriju prostora i kontakt podatke. Fokus portala je na "
        "prodaji i predstavljanju usluga coworking prostora: Daily Pass, 10 Days, Hot Desk, Dedicated "
        "Desk, sale za sastanke, event prostor i poslovna adresa. Kada korisnik kupuje paket ili uslugu, "
        "sistem kreira porudžbinu i pokreće Stripe Checkout ili demo režim plaćanja. Dodatno, korisnik "
        "može da rezerviše meeting room ili phone booth u intervalima od 30 minuta, uz proveru zauzetosti.",
        first_line=True,
    )
    add_paragraph(
        doc,
        "Admin nakon prijave vidi pregled najvažnijih podataka: članove, pakete, plaćanja, Virtual Office "
        "zahteve, uvezenu bazu članova i operativne rezervacije prostora. Admin može da aktivira ili "
        "deaktivira naloge i resurse, ažurira katalog usluga i proveri status transakcija.",
        first_line=True,
    )
    add_bullets(
        doc,
        [
            "Podaci o korisnicima, prostorijama, rezervacijama, proizvodima, porudžbinama i obaveštenjima čuvaju se u PostgreSQL bazama.",
            "Sistem proverava dostupnost prostorije i pravila rezervacije pre upisa u bazu.",
            "Svaka porudžbina ima status pending ili paid, a potvrda plaćanja može doći kroz Stripe webhook.",
            "Administratorske operacije su zaštićene ulogom admin i JWT tokenom.",
        ],
    )


def add_technologies(doc):
    add_section_title(doc, "2", "Korišćene tehnologije")
    add_subtitle(doc, "2.1", "Frontend tehnologije")
    add_paragraph(
        doc,
        "Frontend je izrađen kao moderna HTML, CSS i JavaScript aplikacija u fajlu "
        "spaceforwork-portal.html. Pošto je cilj projekta bio da brzo nastane funkcionalan i vizuelno "
        "uverljiv portal, izabran je jednostavan pristup bez dodatnog frontend build procesa. Aplikacija "
        "ipak radi kao mali SPA: promena stranica, korisnički portal, admin dashboard, katalog paketa, "
        "modal za plaćanje i modul za rezervaciju prostora prikazuju se bez prelaska na odvojene HTML fajlove.",
        first_line=True,
    )
    add_paragraph(
        doc,
        "Na frontend delu su implementirani validacija formi, pregled coworking paketa, kupovina usluga, "
        "prikaz porudžbina, Virtual Office zahtevi, administratorski panel i podrška za rezervaciju "
        "prostora. Dizajn je prilagođen coworking prostoru: koristi realne fotografije Space For Work "
        "prostora, jasnu navigaciju, kartice za pakete i tabove za admin kontrole.",
        first_line=True,
    )

    add_subtitle(doc, "2.2", "Backend tehnologije")
    add_paragraph(
        doc,
        "Backend je razvijen u Node.js okruženju korišćenjem NestJS frameworka i TypeScript jezika. "
        "Aplikacija je podeljena na API Gateway i više mikroservisa, što odgovara zahtevu da sistem "
        "bude lak za održavanje i nadogradnju. API Gateway je jedina ulazna tačka za klijenta, proverava "
        "JWT token i prosleđuje zahteve odgovarajućem servisu.",
        first_line=True,
    )
    add_bullets(
        doc,
        [
            "User Service - registracija, login, JWT tokeni, uloge i admin upravljanje korisnicima.",
            "Room Service - CRUD nad prostorijama, filtriranje, paginacija i SSE tok dostupnosti.",
            "Booking Service - kreiranje, pregled i otkazivanje rezervacija uz poslovna pravila.",
            "Notification Service - čuvanje i isporuka obaveštenja, uključujući RabbitMQ događaje.",
            "Commerce Service - proizvodi, paketi, porudžbine, Stripe Checkout i webhook obrada.",
            "API Gateway - centralizovano rutiranje, autentifikacija i prosleđivanje korisničkog konteksta.",
        ],
    )

    add_subtitle(doc, "2.3", "Baza podataka")
    add_paragraph(
        doc,
        "Za skladištenje podataka koristi se PostgreSQL. Pristup bazi je code-first, kroz TypeORM "
        "entitete i migracije. To znači da su tabele opisane u kodu, a struktura baze prati modele kao "
        "što su User, Room, Booking, Product, Order, OrderLine i Notification. Ovaj pristup je praktičan "
        "za mikroservisnu arhitekturu jer svaki servis može imati svoju oblast podataka i jasnu odgovornost.",
        first_line=True,
    )

    add_subtitle(doc, "2.4", "Razvojna okruženja i alati")
    add_paragraph(
        doc,
        "Za razvoj su korišćeni Visual Studio Code, Git i GitHub. API je testiran preko Swagger-a, "
        "curl komandi i kroz sam portal. Projekat ima Docker Compose konfiguraciju za PostgreSQL, "
        "RabbitMQ, Prometheus i Grafana, kao i Jest testove za ključne servise. Prometheus prikuplja "
        "metrike, dok Grafana služi za pregled stanja sistema.",
        first_line=True,
    )

    add_subtitle(doc, "2.5", "Stripe integracija")
    add_paragraph(
        doc,
        "Stripe je dodat zbog zahteva za online plaćanje. Korišćen je zvanični Stripe SDK za Node.js. "
        "Commerce servis kreira Checkout sesiju, čuva porudžbinu sa statusom pending i nakon Stripe "
        "webhook događaja checkout.session.completed menja status u paid. Ako ključevi nisu podešeni, "
        "aplikacija i dalje može da demonstrira tok plaćanja kroz demo režim, bez unošenja realne kartice.",
        first_line=True,
    )


def add_uml(doc):
    add_section_title(doc, "3", "UML DIJAGRAMI")
    add_subtitle(doc, "3.1", "Dijagram slučajeva upotrebe")
    add_paragraph(
        doc,
        "Dijagram slučajeva upotrebe prikazuje šta korisnik, admin i Stripe sistem rade u okviru "
        "aplikacije. U centru je Space For Work sistem, dok su oko njega akteri koji pokreću glavne "
        "procese: pregled ponude, rezervaciju termina, kupovinu paketa i administraciju.",
        first_line=True,
    )
    add_image(doc, DIAGRAM_DIR / "doc-usecase.png", 12.2, "Slika 1. Dijagram slučajeva upotrebe")

    add_minor_title(doc, "3.1.1", "Zajednički slučajevi upotrebe")
    add_paragraph(
        doc,
        "Registracija - Novi korisnik unosi ime, email i lozinku. Sistem proverava validnost podataka, "
        "jedinstvenost email adrese i čuva nalog sa podrazumevanom member ulogom.",
        first_line=True,
    )
    add_paragraph(
        doc,
        "Prijava - Korisnik ili admin unose email i lozinku. Ako su kredencijali ispravni, User servis "
        "vraća JWT token. Na osnovu uloge korisnik se usmerava na korisnički portal ili admin dashboard.",
        first_line=True,
    )
    add_paragraph(
        doc,
        "Pregled ponude - Svi posetioci mogu da pregledaju coworking opcije, sale, Virtual Office, "
        "private office i pakete. Ovaj deo je važan jer sistem nije samo interna aplikacija, već i javni "
        "prodajni portal prostora.",
        first_line=True,
    )

    add_minor_title(doc, "3.1.2", "Korisnik - slučajevi upotrebe")
    add_paragraph(
        doc,
        "Rezerviši termin - Korisnik bira prostoriju, datum, vreme početka i vreme završetka. Sistem "
        "proverava da li je termin slobodan, da li traje najmanje 30 minuta i da li korisnik nije prešao "
        "limit od tri aktivne rezervacije dnevno.",
        first_line=True,
    )
    add_paragraph(
        doc,
        "Kupi paket - Korisnik bira coworking ili drugu uslugu, sistem formira porudžbinu i pokreće "
        "Stripe Checkout. Ako je plaćanje uspešno, porudžbina prelazi u status paid.",
        first_line=True,
    )
    add_paragraph(
        doc,
        "Pregled istorije - Korisnik može da vidi svoje rezervacije, porudžbine i statuse plaćanja, što "
        "smanjuje potrebu za dodatnim porukama sa administracijom.",
        first_line=True,
    )

    add_minor_title(doc, "3.1.3", "Admin - slučajevi upotrebe")
    add_paragraph(
        doc,
        "Admin upravlja korisnicima, prostorijama, rezervacijama, paketima i transakcijama. U admin panelu "
        "postoje pregledi sa pretragom i filtriranjem, a zaštićene rute zahtevaju admin ulogu. Admin može "
        "da deaktivira korisnika ili prostoriju bez fizičkog brisanja iz baze, čime se čuva istorija sistema.",
        first_line=True,
    )
    add_paragraph(
        doc,
        "Pregled transakcija je posebno važan za zadatak jer admin vidi porudžbinu, stavke, ukupnu cenu, "
        "status, Stripe session ili payment intent referencu i poslednje četiri cifre testne kartice kada "
        "su dostupne.",
        first_line=True,
    )

    add_minor_title(doc, "3.1.4", "Stripe - eksterni sistem")
    add_paragraph(
        doc,
        "Stripe učestvuje u toku plaćanja kao eksterni servis. Aplikacija kreira Checkout sesiju, korisnik "
        "plaća test karticom, a Stripe zatim šalje webhook. Commerce servis proverava potpis webhook-a i "
        "tek nakon validne potvrde menja status porudžbine.",
        first_line=True,
    )

    add_subtitle(doc, "3.2", "Dijagram klasa")
    add_paragraph(
        doc,
        "Dijagram klasa prikazuje glavne modele sistema i njihove veze. Pošto je aplikacija podeljena na "
        "mikroservise, neke veze se čuvaju preko identifikatora, a ne kroz klasične foreign key relacije "
        "između različitih servisa. To je namerna odluka jer servisi ostaju nezavisni.",
        first_line=True,
    )
    add_image(doc, DIAGRAM_DIR / "doc-class.png", 11.3, "Slika 2. Dijagram klasa")

    add_minor_title(doc, "3.2.1", "Opis klasa i atributa")
    add_paragraph(
        doc,
        "User - predstavlja korisnika sistema. Čuva ime, email, hash lozinke, ulogu, aktivnost naloga i "
        "datume kreiranja i izmene. Uloge koje su važne za projekat su admin i member.",
        first_line=True,
    )
    add_paragraph(
        doc,
        "Room - predstavlja prostor koji se može rezervisati. Sadrži naziv, tip, kapacitet, opremu, "
        "aktivnost i datum kreiranja. Tipovi mogu biti sale za sastanke, phone booth ili drugi prostori "
        "koje admin doda kroz panel.",
        first_line=True,
    )
    add_paragraph(
        doc,
        "Booking - predstavlja rezervaciju korisnika. Čuva userId, roomId, datum, vreme početka, vreme "
        "završetka, status i vreme kreiranja. Status omogućava otkazivanje bez brisanja istorije.",
        first_line=True,
    )
    add_paragraph(
        doc,
        "Product - predstavlja uslugu ili paket koji se prodaje, na primer dnevni coworking, mesečni "
        "membership, Virtual Office ili sala za događaje. Čuva cenu, kategoriju, zalihe, opis i oznake.",
        first_line=True,
    )
    add_paragraph(
        doc,
        "Order i OrderLine - predstavljaju porudžbinu i njene stavke. Stavke čuvaju naziv proizvoda i "
        "cenu u trenutku kupovine, pa kasnija promena cene u katalogu ne menja istoriju plaćanja.",
        first_line=True,
    )
    add_paragraph(
        doc,
        "Notification - čuva obaveštenja za korisnika, tip poruke, tekst, status pročitanosti i datum "
        "kreiranja. Booking servis šalje događaje preko RabbitMQ-a, a Notification servis ih obrađuje.",
        first_line=True,
    )

    add_minor_title(doc, "3.2.2", "Opis relacija između klasa")
    add_bullets(
        doc,
        [
            "User -> Booking (1..*) - jedan korisnik može imati više rezervacija, dok svaka rezervacija pripada jednom korisniku preko userId vrednosti.",
            "Room -> Booking (1..*) - jedna prostorija može imati više rezervacija u različitim terminima.",
            "User -> Order (1..*) - jedan korisnik može imati više porudžbina.",
            "Order -> OrderLine (1..*) - porudžbina ima jednu ili više stavki, a brisanjem porudžbine brišu se i njene stavke.",
            "OrderLine -> Product (*..1) - više stavki porudžbina može referencirati isti proizvod.",
            "User -> Notification (1..*) - korisnik može dobiti više obaveštenja.",
        ],
    )


def add_database(doc):
    add_section_title(doc, "4", "Baza podataka")
    add_subtitle(doc, "4.1", "Pristup bazi podataka")
    add_paragraph(
        doc,
        "U projektu je korišćen code-first pristup. Struktura baze se ne piše ručno kao početni SQL "
        "dijagram, već nastaje iz TypeORM entiteta u svakom servisu. Ovakav pristup je pogodan za "
        "NestJS i mikroservise jer se model, validacija i poslovna logika razvijaju zajedno.",
        first_line=True,
    )
    add_paragraph(
        doc,
        "Pošto se koristi code-first pristup, trigger nije centralni deo rešenja. Pravila kao što su "
        "provera preklapanja rezervacija, smanjenje stanja nakon plaćanja i zabrana kupovine preko "
        "dostupne količine implementirana su u servisnom sloju. Time je logika vidljiva u kodu i lakša "
        "za testiranje.",
        first_line=True,
    )

    add_subtitle(doc, "4.2", "Tabele baze podataka")
    add_paragraph(doc, "Baza se sastoji od sledećih glavnih tabela:", bold=True)
    add_bullets(
        doc,
        [
            "users - čuva naloge, email, hash lozinke, uloge, aktivnost naloga i datume izmene.",
            "rooms - čuva prostorije i resurse za rezervaciju, njihov tip, kapacitet, opremu i dostupnost.",
            "bookings - čuva rezervacije, korisnika, prostoriju, datum, vreme početka, vreme završetka i status.",
            "notifications - čuva sistemska obaveštenja nastala iz booking događaja.",
            "products - čuva usluge i pakete koji se prodaju kroz portal.",
            "orders - čuva porudžbine, ukupan iznos, valutu, status i Stripe reference.",
            "order_lines - čuva stavke porudžbine, količinu, cenu po jedinici i ukupnu cenu stavke.",
        ],
    )

    add_subtitle(doc, "4.3", "Poslovna pravila u bazi i servisima")
    add_paragraph(
        doc,
        "Najvažnija pravila ne smeju da zavise od toga da li je korisnik kliknuo ispravno dugme na "
        "frontendu, pa su zato implementirana i u backend servisima. Booking servis proverava da je "
        "vreme završetka posle vremena početka, da rezervacija traje najmanje 30 minuta, da su termini "
        "u koracima od 30 minuta i da ne postoji preklapanje za istu prostoriju ili za istog korisnika.",
        first_line=True,
    )
    add_paragraph(
        doc,
        "Commerce servis proverava da proizvod postoji, da je aktivan i da tražena količina ne prelazi "
        "stanje. Nakon uspešnog plaćanja smanjuje se stanje proizvoda, a porudžbina se označava kao "
        "plaćena. Ovo je važno zato što admin panel kasnije prikazuje realno stanje prodaje i transakcija.",
        first_line=True,
    )


def add_solution(doc):
    add_section_title(doc, "5", "Opis predloženog rešenja")
    add_subtitle(doc, "5.1", "Pregled sistema")
    add_paragraph(
        doc,
        "Predloženo rešenje je web aplikacija koja pokriva ceo osnovni tok rada jednog coworking "
        "prostora: predstavljanje usluga, prodaju paketa i članstva, Virtual Office i private office "
        "ponudu, korisnički portal, administratorsku kontrolu i podršku za rezervaciju prostora. "
        "Rešenje je namerno napravljeno tako da već sada može da se demonstrira kao funkcionalan "
        "coworking informacioni sistem.",
        first_line=True,
    )
    add_image(doc, DIAGRAM_DIR / "04-architecture.png", 15.2, "Slika 3. Arhitektura aplikacije")

    add_subtitle(doc, "5.2", "Korisnički deo aplikacije — coworking portal")
    add_paragraph(
        doc,
        "Početna strana je osmišljena kao prodajni portal za Space For Work. Korisnik dobija jasnu "
        "sliku šta prostor nudi: coworking, membership, Virtual Office, private office, sale za "
        "sastanke, event prostor i kontakt. Portal je vizuelno usklađen sa brendom prostora i služi "
        "kao glavna digitalna prezentacija coworking ponude.",
        first_line=True,
    )
    add_image(
        doc,
        SCREENSHOT_DIR / "portal-home.png",
        15.5,
        "Slika 4. Početna strana Space For Work portala",
    )
    add_paragraph(
        doc,
        "Stranica Coworking predstavlja open space zonu, phone booth-ove, kapacitet prostora i "
        "način rada u coworking okruženju. Membership stranica prikazuje katalog paketa sa cenama "
        "i jasnim opisom šta korisnik dobija za Daily Pass, 10 Days, Hot Desk i Dedicated Desk.",
        first_line=True,
    )
    add_image(
        doc,
        SCREENSHOT_DIR / "portal-coworking.png",
        15.5,
        "Slika 5. Coworking stranica portala",
    )
    add_image(
        doc,
        SCREENSHOT_DIR / "portal-membership.png",
        15.5,
        "Slika 6. Membership paketi i katalog usluga",
    )
    add_paragraph(
        doc,
        "Korisnički portal prikazuje kupljene pakete, porudžbine, Virtual Office zahteve i, po "
        "potrebi, rezervacije prostora. Korisnik ne mora da pamti šta je poslao ili platio, jer se "
        "istorija čuva u sistemu.",
        first_line=True,
    )

    add_subtitle(doc, "5.3", "Modul za rezervaciju prostora")
    add_paragraph(
        doc,
        "Pored prodaje coworking usluga, sistem podržava i rezervaciju meeting room-ova, phone "
        "booth-ova i event prostora. Slobodni i zauzeti intervali su vizuelno razdvojeni, a forma "
        "vodi korisnika kroz rezervaciju bez nepotrebnih koraka. Ovaj modul je dopuna glavnom "
        "coworking sistemu, a ne njegova centralna tačka.",
        first_line=True,
    )
    add_subtitle(doc, "5.4", "Administratorski panel")
    add_paragraph(
        doc,
        "Admin panel je centralno mesto za upravljanje coworking poslovanjem. Nakon prijave admin "
        "vidi kartice sa ključnim pokazateljima, članove, pakete, transakcije, Virtual Office zahteve, "
        "uvezenu bazu članova i operativne rezervacije prostora.",
        first_line=True,
    )
    add_image(
        doc,
        SCREENSHOT_DIR / "portal-admin.png",
        15.5,
        "Slika 7. Administratorska kontrolna tabla",
    )
    add_bullets(
        doc,
        [
            "Upravljanje korisnicima i članstvima - pregled, pretraga, izmena uloge i aktivacije naloga.",
            "Upravljanje katalogom paketa - dodavanje i ažuriranje coworking usluga koje se mogu kupiti.",
            "Pregled transakcija - status plaćanja, Stripe reference, stavke porudžbine i iznosi.",
            "Virtual Office i private office - praćenje zahteva, ugovora i aktivnih korisnika usluge.",
            "Upravljanje prostorijama i rezervacijama - dodavanje resursa i pregled termina po potrebi.",
            "Pregled Office Members baze - korišćenje postojeće coworking baze kao poslovnog konteksta.",
        ],
    )

    add_subtitle(doc, "5.5", "Autentifikacija i autorizacija")
    add_paragraph(
        doc,
        "Sistem koristi JWT autentifikaciju. Nakon uspešnog logina User servis vraća token koji sadrži "
        "ID korisnika, email i ulogu. API Gateway proverava token i prosleđuje x-user-id i x-user-role "
        "ka mikroservisima. Na taj način mikroservisi znaju ko šalje zahtev, a za admin rute mogu da "
        "odbiju korisnika koji nema odgovarajuću ulogu.",
        first_line=True,
    )
    add_paragraph(
        doc,
        "Demo nalozi koji se koriste za prikaz sistema su admin@spaceforwork.rs sa lozinkom admin123 i "
        "korisnik@spaceforwork.rs sa lozinkom korisnik123. Prvi nalog ima administratorsku ulogu, dok "
        "drugi predstavlja običnog člana prostora.",
        first_line=True,
    )

    add_subtitle(doc, "5.6", "REST API i integracije")
    add_paragraph(
        doc,
        "API je organizovan tako da javne rute mogu da koriste i gosti, dok zaštićene rute zahtevaju "
        "JWT token. Javno su dostupni registracija, login, pregled prostorija, zauzetost prostorije po "
        "datumu i pregled paketa. Zaštićene rute omogućavaju kreiranje rezervacije, pregled mojih "
        "rezervacija, kupovinu paketa, pregled mojih porudžbina i admin operacije.",
        first_line=True,
    )
    add_paragraph(
        doc,
        "Između servisa se koriste tri tipa komunikacije. REST se koristi za standardne zahteve preko "
        "Gateway-a, RabbitMQ za slanje događaja iz Booking servisa ka Notification servisu, a SSE za "
        "reaktivne prikaze obaveštenja i dostupnosti.",
        first_line=True,
    )

    add_subtitle(doc, "5.7", "Stripe plaćanje i webhook")
    add_paragraph(
        doc,
        "Plaćanje je implementirano kroz Commerce servis. Kada korisnik kupi paket, servis proverava "
        "proizvode i količine, kreira porudžbinu sa statusom pending i zatim kreira Stripe Checkout sesiju. "
        "Korisnik se preusmerava na Stripe stranicu za testno plaćanje. Za demonstraciju se koristi test "
        "kartica, na primer 4242 4242 4242 4242.",
        first_line=True,
    )
    add_paragraph(
        doc,
        "Nakon uspešnog plaćanja Stripe šalje webhook na endpoint /webhooks/stripe. Commerce servis "
        "proverava Stripe-Signature header i signing secret. Tek ako je potpis ispravan i događaj je "
        "checkout.session.completed, porudžbina prelazi u status paid, čuva se payment intent referenca "
        "i smanjuje se stanje proizvoda.",
        first_line=True,
    )
    add_paragraph(
        doc,
        "Ako Stripe ključevi nisu podešeni, aplikacija nastavlja da radi i ne traži realnu karticu. Tada "
        "se koristi demo režim plaćanja, što je praktično za lokalnu prezentaciju projekta i za situacije "
        "kada nije potrebno povezivati pravi Stripe nalog.",
        first_line=True,
    )
    add_image(doc, DIAGRAM_DIR / "05-payment-extend.png", 14.2, "Slika 8. Tok plaćanja i webhook potvrde")

    add_subtitle(doc, "5.8", "Validacija i obrada grešaka")
    add_paragraph(
        doc,
        "Validacija postoji na više nivoa. Frontend proverava da korisnik ne pošalje prazna ili očigledno "
        "neispravna polja, dok backend ima DTO validaciju i poslovna pravila. Na primer, booking servis "
        "vraća grešku ako je termin kraći od 30 minuta, ako nije u pravilnom intervalu ili ako se preklapa "
        "sa postojećom rezervacijom. Commerce servis vraća grešku ako se pokuša kupovina nedostupnog "
        "proizvoda ili količine koja prelazi stanje.",
        first_line=True,
    )
    add_paragraph(
        doc,
        "Greške se obrađuju kroz NestJS exception mehanizam. To znači da aplikacija korisniku ne prikazuje "
        "nejasan pad sistema, već poruku koja objašnjava šta treba popraviti. U admin delu to je posebno "
        "korisno jer osoba koja upravlja prostorom odmah vidi zbog čega određena akcija nije prošla.",
        first_line=True,
    )

    add_subtitle(doc, "5.9", "Povezanost sa zahtevima projektnog zadatka")
    add_paragraph(
        doc,
        "U odnosu na zahteve predmeta, projekat pokriva autentifikaciju, role, korisnički i admin deo, "
        "CRUD operacije, pretragu i filtriranje, validaciju, rad sa bazom, online plaćanje, webhook, "
        "pregled transakcija i UML dokumentaciju. Sistem ima dovoljno poslovne logike da ne bude samo "
        "vizuelni prototip, ali je zadržan dovoljno jednostavno da se može jasno objasniti i odbraniti.",
        first_line=True,
    )


def add_conclusion(doc):
    add_section_title(doc, "6", "Zaključak")
    add_paragraph(
        doc,
        "U okviru projekta razvijen je informacioni sistem za Space For Work coworking prostor. Sistem "
        "objedinjuje javni coworking sajt, prodaju paketa i članstva, korisnički portal, admin dashboard, "
        "bazu podataka, mikroservisni backend, testno online plaćanje i modul za rezervaciju prostora. "
        "Time je pokriven realan problem coworking prostora: kako korisnicima brzo prikazati usluge, "
        "omogućiti kupovinu paketa i Virtual Office ponude, a administratoru dati pregled nad "
        "svakodnevnim poslovanjem.",
        first_line=True,
    )
    add_paragraph(
        doc,
        "Najvažniji deo rešenja je to što su ključna pravila implementirana u backendu, a ne samo u "
        "interfejsu. Sistem ne dozvoljava preklapanje rezervacija, ograničava broj aktivnih termina po "
        "korisniku, proverava stanje paketa pre kupovine i menja status porudžbine tek nakon Stripe "
        "potvrde ili kontrolisanog demo plaćanja.",
        first_line=True,
    )
    add_paragraph(doc, "Urađeno je:", bold=True, space_before=6)
    add_bullets(
        doc,
        [
            "javni Space For Work portal sa coworking, membership, Virtual Office i private office stranicama,",
            "katalog paketa i online kupovina coworking usluga,",
            "login za admina i korisnika,",
            "admin dashboard za kontrolu članova, paketa, transakcija i Virtual Office zahteva,",
            "modul za rezervaciju meeting room-ova i phone booth-ova,",
            "PostgreSQL baze sa TypeORM entitetima,",
            "Stripe Checkout, webhook i demo režim plaćanja,",
            "UML dijagram slučajeva upotrebe i dijagram klasa,",
            "dokumentacija u formatu koji prati zahtev predmeta.",
        ],
    )
    add_paragraph(
        doc,
        "Dalja nadogradnja može da obuhvati potpunu produkcionu Stripe konfiguraciju, slanje email "
        "potvrda, napredniji kalendar za administratore, automatsko fakturisanje, loyalty sistem za "
        "članove i bolju analitiku popunjenosti prostora. Osnovna arhitektura je već pripremljena tako "
        "da ove nadogradnje ne zahtevaju promenu cele aplikacije, već proširenje postojećih servisa.",
        first_line=True,
    )
    add_blank(doc, 2)
    add_paragraph(doc, "Linkovi projekta:", bold=True, space_after=2)
    p = doc.add_paragraph()
    add_hyperlink(p, "GitHub repozitorijum", REPO_URL)
    p2 = doc.add_paragraph()
    add_hyperlink(p2, "Demo portal (lokalno)", PORTAL_URL)


def build():
    doc = configure_document()
    add_cover(doc)
    add_contents(doc)
    add_real_system(doc)
    add_technologies(doc)
    add_uml(doc)
    add_database(doc)
    add_solution(doc)
    add_conclusion(doc)
    OUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUT)
    print(f"Generated {OUT}")


if __name__ == "__main__":
    build()
